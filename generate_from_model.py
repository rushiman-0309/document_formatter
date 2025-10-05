import fitz  # PyMuPDF
from docx import Document
import os
import json
import google.generativeai as genai
from dotenv import load_dotenv
load_dotenv()
# --- 1. Gemini API Configuration ---
try:
    api_key = os.getenv("GOOGLE_API_KEY")
    # api_key = 
    
    if not api_key:
        raise ValueError(
            "API key not found. Please set the GOOGLE_API_KEY environment variable.")
    genai.configure(api_key=api_key)
    print("Gemini API configured successfully.")
except (ValueError, Exception) as e:
    print(f"Error configuring Gemini API: {e}")
    exit()

# --- 2. Core Functions ---


def extract_text_from_docx(docx_path):
    """Extracts all text content from a .docx file."""
    try:
        doc = Document(docx_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        return '\n'.join(full_text)
    except Exception as e:
        print(f"Error reading DOCX file: {e}")
        return ""


def infer_placeholders_with_gemini(template_text):
    """
    Analyzes the text of a model document and asks Gemini to identify the
    variable fields and their example values.
    """
    model = genai.GenerativeModel('gemini-2.0-flash')
    prompt = f"""
    You are a template analysis expert. Analyze the following text from a model document.
    Your task is to identify all the pieces of information that are specific to the example and would change in a different report.
    For each piece of variable data, create a logical, snake_case key name (e.g., 'product_name', 'nfpa_health_rating').
    Return a single JSON object where the keys are your inferred key names and the values are the corresponding example text from the document.
    For the 'nfpa_diamond', represent it as a multi-line string.

    MODEL DOCUMENT TEXT:
    ---
    {template_text}
    ---
    """
    print("Phase 1: Analyzing model document with Gemini to infer structure...")
    try:
        response = model.generate_content(prompt)
        json_text = response.text.strip().lstrip("```json").rstrip("```").strip()
        inferred_data = json.loads(json_text)
        print("...Analysis complete. Inferred structure:")
        print(json.dumps(inferred_data, indent=2))
        return inferred_data
    except Exception as e:
        print(f"Error during template analysis: {e}")
        return {}


def extract_text_from_pdf(pdf_path):
    """Extracts all text from a PDF file."""
    try:
        with fitz.open(pdf_path) as doc:
            return "".join(page.get_text() for page in doc)
    except Exception as e:
        print(f"Error reading PDF file: {e}")
        return None


def get_data_from_gemini(fields_to_find, pdf_text):
    """Asks Gemini to find data for a dynamic list of fields from the PDF text."""
    model = genai.GenerativeModel('gemini-2.0-flash')
    fields_list = "\n".join(f"- {field}" for field in fields_to_find)
    prompt = f"""
    You are an expert data extraction system for Safety Data Sheets (SDS).
    Analyze the raw SDS text below. Your task is to find the information for the following fields.
    Return the data ONLY in a valid JSON format where the keys exactly match the field names provided.
    
    FIELDS TO EXTRACT:
    {fields_list}
    
    RAW SDS TEXT:
    ---
    {pdf_text}
    ---
    """
    print("\nPhase 2: Extracting real data from PDF based on inferred structure...")
    try:
        response = model.generate_content(prompt)
        json_text = response.text.strip().lstrip("```json").rstrip("```").strip()
        print("...Extraction complete.")
        return json.loads(json_text)
    except Exception as e:
        print(f"Error during data extraction: {e}")
        return {}


def translate_fields(data_dict, fields_to_translate):
    """Translates specific values in the data dictionary to Hindi."""
    print("\nPhase 3: Translating designated fields...")
    model = genai.GenerativeModel('gemini-2.0-flash')
    for field in fields_to_translate:
        if field in data_dict and data_dict[field]:
            english_text = data_dict[field]
            try:
                prompt = f"Translate the following English phrase to Hindi. Provide only the translation. Phrase: '{english_text}'"
                response = model.generate_content(prompt)
                translated_text = response.text.strip()
                data_dict[field] = translated_text
                print(f"  - Translated '{field}'")
            except Exception:
                print(f"  - Warning: Translation failed for field '{field}'.")
    return data_dict


def docx_replace(doc, replacements):
    """Performs a find-and-replace for all text in a document."""
    for old_text, new_text in replacements.items():
        # Replace in paragraphs
        for p in doc.paragraphs:
            if old_text in p.text:
                inline = p.runs
                for i in range(len(inline)):
                    if old_text in inline[i].text:
                        text = inline[i].text.replace(
                            str(old_text), str(new_text))
                        inline[i].text = text
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if old_text in p.text:
                            inline = p.runs
                            for i in range(len(inline)):
                                if old_text in inline[i].text:
                                    text = inline[i].text.replace(
                                        str(old_text), str(new_text))
                                    inline[i].text = text


def generate_final_document(template_path, output_path, template_data, final_data):
    """Fills the model document by replacing example text with final data."""
    print("\nPhase 4: Generating final document by replacing model data...")
    try:
        doc = Document(template_path)
        replacements = {}
        for key, old_value in template_data.items():
            # Check if the key exists in the final data extracted from the PDF
            if key in final_data:
                new_value = final_data[key]
                # Add to our replacement dictionary
                # Important: Convert to string to avoid errors with non-string types
                replacements[str(old_value)] = str(new_value)

        docx_replace(doc, replacements)
        doc.save(output_path)
        print(f"\nSuccessfully created final document: {output_path}")

    except Exception as e:
        print(f"Error while generating the final document: {e}")


# --- 3. Main Execution Block ---
if __name__ == "__main__":
    # --- CONFIGURATION ---
    MODEL_DOCX_PATH = 'Model_Template.docx'
    # <-- IMPORTANT: Change this
    file_name = 'SDS_potassium-hydroxide-solid_original-id-150559.pdf'
    INPUT_PDF_PATH = os.path.join("input_data", file_name)
    output_filename = f"{file_name.split('.')[0]}_converted.docx"
    OUTPUT_DOCX_PATH = os.path.join("output_data", output_filename)
    # Specify which of the INFERRED keys should be translated
    # These must match the keys Gemini creates in the analysis phase
    FIELDS_TO_TRANSLATE = [
        'document_title',           
        'document_subtitle',  # e.g., "Safety Data Sheet (SDS)"
        'pictogram_notes',
        'ppe_header',
        'ppe_notes',
        'hmis_table_header_category',
        'hmis_table_header_rating',
        'hmis_health_category_label',
        'hmis_flammability_category_label',
        'hmis_physical_hazard_category_label',
        'hmis_ppe_category_label'
    ]

    # Phase 1: Analyze the model document to understand its structure
    template_text = extract_text_from_docx(MODEL_DOCX_PATH)
    if template_text:
        inferred_template_data = infer_placeholders_with_gemini(template_text)

        if inferred_template_data:
            # Phase 2: Extract real data from the PDF using the inferred structure
            pdf_text = extract_text_from_pdf(INPUT_PDF_PATH)
            if pdf_text:
                final_extracted_data = get_data_from_gemini(
                    inferred_template_data.keys(), pdf_text)

                if final_extracted_data:
                    # Phase 3: Translate the designated fields
                    final_data_translated = translate_fields(
                        final_extracted_data, FIELDS_TO_TRANSLATE)

                    # Phase 4: Generate the final report
                    generate_final_document(
                        MODEL_DOCX_PATH, OUTPUT_DOCX_PATH, inferred_template_data, final_data_translated)
